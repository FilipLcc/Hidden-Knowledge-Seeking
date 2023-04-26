import aspose.words as aw
import mysql.connector
import docx

#create a word document
#https://blog.aspose.com/2021/10/28/create-word-documents-using-python/
#NOVI LINK ZA BESPLATAN API
#https://www.androidauthority.com/how-to-write-to-a-file-in-python-1141195/
#TRECI LINK, ODAVDE SI SKINUO
#https://stackabuse.com/reading-and-writing-ms-word-files-in-python-via-python-docx-module/

# C:\Users\Luka\Desktop\PSZ bolje ".\Results.docx"
mydoc = docx.Document()
#connect to the db
mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        database = "pszproject"
        #password="mypassword"
    )
mycursor = mydb.cursor()

#zahtev 1
mydoc.add_paragraph("Rezultati")
para = mydoc.add_paragraph("Broj automobila po markama\n\n")

mycursor.execute("select DISTINCT marka FROM big_kola")
distinct_brands = mycursor.fetchall()
distinct_brands_count = []
distinct_brand_names = []
for brand in distinct_brands:
    #sql = "SELECT count(*) FROM big_kola where marka = %s"
    #had problems with injection defense, so I will merge strings instead
    sql = "SELECT count(*) FROM big_kola where marka = '" + brand[0] + "'"
    #val = {"marka":brand[0]}
    #print(val)
    #mycursor.execute(sql, val)
    mycursor.execute(sql)
    num_of_cars = mycursor.fetchone()
    distinct_brands_count.append(num_of_cars[0])
    distinct_brand_names.append(brand[0])

for i in range (len(distinct_brand_names)):
    #write them to file
    # create font
    para.add_run(distinct_brand_names[i] +": "+ str(distinct_brands_count[i]) +"\n")

#zahtev 2
para = mydoc.add_paragraph("Broj automobila po lokacijama\n\n")


mycursor.execute("select DISTINCT lokacija FROM big_kola")
distinct_locations = mycursor.fetchall()
distinct_location_names = []
cars_per_location = []
for location in distinct_locations:
    if(location[0]==""):
        distinct_location_names.append("Nepoznato")
    else:
        distinct_location_names.append(location[0])
    sql = "SELECT count(*) FROM big_kola where lokacija = '" + location[0] + "'"
    mycursor.execute(sql)
    num_of_cars = mycursor.fetchone()
    cars_per_location.append(num_of_cars[0])

# create heading

for i in range(len(distinct_locations)):
    para.add_run(distinct_location_names[i] +": "+ str(cars_per_location[i]) +"\n")
    

#zahtev 3
para = mydoc.add_paragraph("Broj automobila po bojama\n\n")

mycursor.execute("select distinct boja from big_kola")

distinct_colors = mycursor.fetchall()
distinct_colors_count = []
distinct_color_names = []
for color in distinct_colors:
    sql = "SELECT count(*) FROM big_kola where boja = '" + color[0] + "'"
    mycursor.execute(sql)
    num_of_cars = mycursor.fetchone()
    distinct_colors_count.append(num_of_cars[0])
    if(color[0]==''):
        distinct_color_names.append('nepoznata boja')
    else:
        distinct_color_names.append(color[0])

# create heading

for i in range (len(distinct_color_names)):
        para.add_run(distinct_color_names[i] +": "+ str(distinct_colors_count[i]) +"\n")


#zahtev 4
para = mydoc.add_paragraph("Najskuplji automobili\n\n")

mycursor.execute("SELECT id, marka, model, godiste, lokacija, cena from big_kola order by cena desc limit 30;")
most_expensive_cars = mycursor.fetchall()
mycursor.execute("SELECT id, marka, model, godiste, lokacija, cena, karoserija from big_kola where karoserija like '%SUV%' order by cena desc limit 30")
most_expensive_suvs = mycursor.fetchall()
for i in range (len(most_expensive_cars)):
    #write them to file
    if most_expensive_cars[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = most_expensive_cars[i][4]
    para.add_run( str(i+1) +": "+str(most_expensive_cars[i][0]) +", "+most_expensive_cars[i][1] +", "+most_expensive_cars[i][2]+
    ", " + str(most_expensive_cars[i][3]) +", " + lokacija +", " + str(most_expensive_cars[i][5]) +"\n")

para = mydoc.add_paragraph("Najskuplji iz SUV kategorije\n\n")
for i in range (len(most_expensive_suvs)):
    #write them to file
    #print(most_expensive_suvs[i])
    if most_expensive_suvs[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = most_expensive_suvs[i][4]
    para.add_run(str(i+1) +": "+str(most_expensive_suvs[i][0]) +", "+most_expensive_suvs[i][1] +", "+most_expensive_suvs[i][2]+
    ", " + str(most_expensive_suvs[i][3]) +", " + lokacija +", " + str(most_expensive_suvs[i][5]) +"\n")

#zahtev 5
para = mydoc.add_paragraph("Najnoviji automobili\n\n")

mycursor.execute("SELECT id, marka, model, godiste, lokacija, cena from big_kola where godiste = 2021 or godiste = 2022 order by cena desc")
newest_cars = mycursor.fetchall()
for i in range(len(newest_cars)):
    #write them to files
    if newest_cars[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = newest_cars[i][4]
    para.add_run(str(i+1) +": "+str(newest_cars[i][0]) +", "+newest_cars[i][1] +", "+newest_cars[i][2]+
    ", " + str(newest_cars[i][3]) +", " + lokacija +", " + str(newest_cars[i][5]) +"\n")

#zahtev 6 

#a
para = mydoc.add_paragraph("Automobili sa maksimalnom cenom\n\n")

mycursor.execute("SELECT MAX(cena) FROM big_kola")
max_price = mycursor.fetchone()
sql = "SELECT id, marka, model, godiste, lokacija, cena FROM big_kola where cena = %s"
mycursor.execute(sql, (max_price))
peaking_cars = mycursor.fetchall()
for i in range(len(peaking_cars)):
    #write them to files
    if peaking_cars[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = peaking_cars[i][4]
    para.add_run(str(i+1) +": "+str(peaking_cars[i][0]) +", "+peaking_cars[i][1] +", "+peaking_cars[i][2]+
    ", " + str(peaking_cars[i][3]) +", " + lokacija +", " + str(peaking_cars[i][5]) +"\n")

#b
para = mydoc.add_paragraph("Automobili sa maksimalnom konjskom snagom\n\n")

mycursor.execute("SELECT MAX(snaga) FROM big_kola")
max_power = mycursor.fetchone()
sql = "SELECT id, marka, model, godiste, lokacija, cena, snaga FROM big_kola where snaga = %s"
mycursor.execute(sql, (max_power))
most_powerful = mycursor.fetchall()
for i in range(len(most_powerful)):
    #write them to files
    if most_powerful[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = most_powerful[i][4]
    para.add_run(str(i+1) +": "+str(most_powerful[i][0]) +", "+most_powerful[i][1] +", "+most_powerful[i][2]+
    ", " + str(most_powerful[i][3]) +", " + lokacija +", " + str(most_powerful[i][5]) +"\n")

#c
para = mydoc.add_paragraph("Automobili sa najvecom predjenom kilometrazom\n\n")

mycursor.execute("SELECT MAX(kilometraza) FROM big_kola")
most_km_driven = mycursor.fetchone()
sql = "SELECT id, marka, model, godiste, lokacija, cena, kilometraza FROM big_kola where kilometraza = %s"
mycursor.execute(sql, (most_km_driven))
most_driven = mycursor.fetchall()
for i in range(len(most_driven)):
    #write them to files
    if most_driven[i][4] == "":
        lokacija ="Nepoznato"
    else:
        lokacija = most_driven[i][4]
    para.add_run(str(i+1) +": "+str(most_driven[i][0]) +", "+most_driven[i][1] +", "+most_driven[i][2]+
    ", " + str(most_driven[i][3]) +", " + lokacija +", " + str(most_driven[i][5]) + ", " + str(most_driven[i][6])+"\n")




mydoc.save(".\Results.docx")






